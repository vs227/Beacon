import os
import sys
import shutil
import json
from pathlib import Path

# Add Backend root to path
backend_dir = Path(__file__).resolve().parent.parent.parent
if str(backend_dir) not in sys.path:
    sys.path.insert(0, str(backend_dir))

from RAG.dataIngestion.loader import process_all_files
from RAG.dataIngestion.chunker import split_documents
from RAG.dataIngestion.embeddings import generate_embeddings


def create_sample_files(test_dir: Path):
    """Create sample files for all 6 supported formats."""
    test_dir.mkdir(parents=True, exist_ok=True)

    # 1. TXT
    (test_dir / "notes.txt").write_text(
        "Beacon is an AI Observability and Fleet Control Platform.\n"
        "It provides real-time monitoring and document RAG capabilities.",
        encoding="utf-8"
    )

    # 2. Markdown
    (test_dir / "README.md").write_text(
        "# Beacon Documentation\n\n"
        "## Architecture\n"
        "Beacon uses FastAPI for backend services, Supabase pgvector for vector storage, "
        "and sentence-transformers all-MiniLM-L6-v2 for CPU embedding generation.",
        encoding="utf-8"
    )

    # 3. CSV
    (test_dir / "employees.csv").write_text(
        "id,name,role,department\n"
        "1,Alice,AI Engineer,Research\n"
        "2,Bob,Fullstack Developer,Engineering\n",
        encoding="utf-8"
    )

    # 4. JSON
    (test_dir / "data.json").write_text(
        json.dumps({
            "project": "Beacon Customer Support AI",
            "version": "1.0.0",
            "features": ["Vector Search", "Document Ingestion", "FastAPI"]
        }, indent=2),
        encoding="utf-8"
    )

    # 5. DOCX (Create a real .docx using docx library)
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Beacon Project Manual", 0)
        doc.add_paragraph("This document outlines the core architecture and ingestion guidelines.")
        doc.save(test_dir / "documentation.docx")
    except Exception as e:
        print(f"Skipped docx sample creation: {e}")


def run_test():
    print("=" * 60)
    print("BEACON RAG DATA INGESTION PIPELINE TEST")
    print("=" * 60)

    test_dir = backend_dir / "RAG" / "dataIngestion" / "test_data"
    create_sample_files(test_dir)

    print(f"\nScanning directory: {test_dir.name}/\n")

    # 1. Load files
    documents = process_all_files(test_dir)
    print(f"\nTotal documents loaded: {len(documents)}")

    # 2. Split into chunks
    chunks = split_documents(documents)

    # 3. Generate embeddings
    texts = [c.page_content for c in chunks]
    embeddings = generate_embeddings(texts, batch_size=32)

    print(f"\nEmbedding shape: {embeddings.shape}")
    print(f"Successfully processed and generated {len(chunks)} vectors (384-D).")
    print("=" * 60)

    # Cleanup test files
    shutil.rmtree(test_dir, ignore_errors=True)


if __name__ == "__main__":
    run_test()
